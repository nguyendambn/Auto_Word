import re
import os
import logging
logging.basicConfig(level=logging.WARNING, format="%(levelname)s: %(message)s")
from lxml import etree
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def add_page_number_field(run):
    """Chèn trường PAGE vào run (đánh số trang động trong Word)."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Giá trị hiển thị tạm trước khi Word cập nhật
    default_val = OxmlElement('w:t')
    default_val.text = "1"

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(default_val)
    run._r.append(fldChar3)


def clear_header_footer(hf):
    """Xóa sạch nội dung (bao gồm paragraphs, tables và các block content control như sdt) của một Header/Footer."""
    for child in list(hf._element):
        try:
            hf._element.remove(child)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
    # Luôn thêm lại một paragraph rỗng mặc định để đảm bảo tính hợp lệ của schema Word
    try:
        hf.add_paragraph()
    except Exception as e:
            logging.warning(f"Silent error ignored: {e}")


def clear_all_headers_footers(doc, skip_sections=None):
    """Xóa sạch mọi header/footer và hủy liên kết (unlink) của tất cả các sections."""
    if skip_sections is None:
        skip_sections = set()
    for idx, section in enumerate(doc.sections):
        if idx in skip_sections:
            continue
        section.header.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        try:
            section.even_page_header.is_linked_to_previous = False
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
            
        section.footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        try:
            section.even_page_footer.is_linked_to_previous = False
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
            
        clear_header_footer(section.header)
        clear_header_footer(section.first_page_header)
        try:
            clear_header_footer(section.even_page_header)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
            
        clear_header_footer(section.footer)
        clear_header_footer(section.first_page_footer)
        try:
            clear_header_footer(section.even_page_footer)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")



def set_run_font(run, font_name, font_size=None, bold=None, italic=None):
    """Thiết lập font chữ cho Run, bao gồm cả hỗ trợ ký tự tiếng Việt (eastAsia, cs)."""
    if font_name:
        run.font.name = font_name
        rPr = run.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
        try: run.font.cs_bold = bold
        except Exception: pass
    if italic is not None:
        run.italic = italic
        try: run.font.cs_italic = italic
        except Exception: pass
    # Thiết lập màu đen cho tất cả chữ trong tài liệu
    run.font.color.rgb = RGBColor(0, 0, 0)

# XML element order schema lists for Word validation
PPR_ORDER = [
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
    'widowControl', 'numPr', 'pBdr', 'shd', 'tabs', 'spacing', 'ind',
    'contextualSpacing', 'mirrorMargins', 'textboxTightWrap', 'jc',
    'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr'
]

TBLPR_ORDER = [
    'tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
    'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
    'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption',
    'tblDescription'
]

SETTINGS_ORDER = [
    'writeProtection', 'view', 'zoom', 'removePersonalInformation', 'removeDateAndTime',
    'doNotDisplayPageBoundaries', 'displayBackgroundShape', 'printPostScriptOverText',
    'printFractionalCharacterWidth', 'printFormsData', 'embedTrueTypeFonts',
    'embedSystemFonts', 'saveSubsetFonts', 'saveFormsData', 'mirrorMargins',
    'alignBordersAndEdges', 'bordersDoNotSurroundHeader', 'bordersDoNotSurroundFooter',
    'gutterAtTop', 'hideSpellingErrors', 'hideGrammaticalErrors', 'activeWritingStyle',
    'proofState', 'formsDesign', 'attachedTemplate', 'linkStyles',
    'stylePaneFormatFilter', 'stylePaneSortMethod', 'documentType', 'mailMerge',
    'revisionView', 'trackRevisions', 'doNotTrackMoves', 'doNotTrackFormatting',
    'documentProtection', 'autoFormatOverride', 'styleLockTheme', 'styleLockQFSet',
    'defaultTabStop', 'autoHyphenation', 'consecutiveHyphenLimit', 'hyphenationZone',
    'doNotHyphenateCaps', 'showEnvelope', 'summaryLength', 'clickAndTypeStyle',
    'defaultTableStyle', 'evenAndOddHeaders', 'bookFoldRevPrinting',
    'bookFoldPrinting', 'bookFoldPrintingSheets', 'drawingGridHorizontalSpacing',
    'drawingGridVerticalSpacing', 'displayHorizontalDrawingGridEvery',
    'displayVerticalDrawingGridEvery', 'doNotUseMarginsForDrawingGridOrigin',
    'drawingGridHorizontalOrigin', 'drawingGridVerticalOrigin',
    'doNotShadeFormData', 'noPunctuationKerning', 'characterSpacingControl',
    'printTwoOnOne', 'strictFirstAndLastChars', 'noLineBreaksAfter',
    'noLineBreaksBefore', 'savePreviewPicture', 'doNotValidateAgainstSchema',
    'saveInvalidXml', 'ignoreMixedContent', 'alwaysShowPlaceholderText',
    'doNotDemarcateInvalidXml', 'saveXmlDataOnly', 'useXSLTWhenSaving',
    'saveThroughXslt', 'showXMLTags', 'alwaysMergeEmptyNamespace',
    'updateFields', 'hdrShapeDefaults', 'footnotePr', 'endnotePr',
    'compat', 'docVars', 'rsids', 'mathPr', 'uiCompat97To2003',
    'attachedSchema', 'themeFontLang', 'clrSchemeMapping',
    'doNotIncludeSubdocsInStats', 'doNotAutoCompressPictures',
    'forceUpgrade', 'captions', 'readModeInkLockDown', 'smartTagType',
    'schemaLibrary', 'shapeDefaults', 'doNotEmbedSmartTags',
    'decimalSymbol', 'listSeparator'
]

SECTPR_ORDER = [
    'headerReference', 'footerReference', 'type', 'pgSz', 'pgMar', 'paperSrc',
    'pgBorders', 'lnNumType', 'pgNumType', 'cols', 'vAlign', 'noEndnote',
    'titlePg', 'textDirection', 'presentationDirection', 'bidi', 'rtlGutter',
    'docGrid', 'printerSettings', 'sectPrChange'
]

def insert_element_in_order(parent, element, order_list):
    """Chèn một thẻ XML vào đúng vị trí quy định của XML Schema để tránh hỏng file."""
    tag_name = element.tag.split('}')[-1]
    if tag_name not in order_list:
        parent.append(element)
        return
    idx = order_list.index(tag_name)
    inserted = False
    for child in parent:
        c_tag = child.tag.split('}')[-1]
        if c_tag in order_list:
            c_idx = order_list.index(c_tag)
            if c_idx > idx:
                child.addprevious(element)
                inserted = True
                break
    if not inserted:
        parent.append(element)

def safe_set_style(doc, obj, style_name, fallback='Normal'):
    """Gán style cho paragraph hoặc run một cách an toàn, nếu style không tồn tại thì fallback về Normal."""
    try:
        obj.style = doc.styles[style_name]
    except KeyError:
        try:
            obj.style = doc.styles[fallback]
        except KeyError:
            try:
                obj.style = 'Normal'
            except Exception as e:
                    logging.warning(f"Silent error ignored: {e}")

def set_contextual_spacing(obj, val=True):
    """Bật/tắt thuộc tính 'Don't add space between paragraphs of the same style' (contextualSpacing)."""
    pPr = obj.element.get_or_add_pPr() if hasattr(obj, 'element') else obj._p.get_or_add_pPr()
    contextualSpacing = pPr.find(qn('w:contextualSpacing'))
    if val:
        if contextualSpacing is None:
            contextualSpacing = OxmlElement('w:contextualSpacing')
            insert_element_in_order(pPr, contextualSpacing, PPR_ORDER)
    else:
        if contextualSpacing is not None:
            pPr.remove(contextualSpacing)


def paragraph_has_image(p):
    """Kiểm tra xem paragraph có chứa hình ảnh hay không (namespace-agnostic)."""
    elements = p._p.xpath('.//*[local-name()="drawing" or local-name()="pict" or local-name()="inline" or local-name()="graphic" or local-name()="imagedata" or local-name()="pic" or local-name()="object"]')
    return len(elements) > 0


def roman_to_int(s):
    s = s.upper()
    roman_map = {'I':1, 'V':5, 'X':10, 'L':50, 'C':100, 'D':500, 'M':1000}
    val = 0
    for i in range(len(s)):
        if i > 0 and roman_map[s[i]] > roman_map[s[i-1]]:
            val += roman_map[s[i]] - 2 * roman_map[s[i-1]]
        else:
            val += roman_map[s[i]]
    return val


def get_chapter_number(text):
    m = re.match(r'^(?:CHƯƠNG|Chương)\s+([IVX\d]+)', text, re.IGNORECASE)
    if m:
        val = m.group(1)
        if val.isdigit():
            return int(val)
        else:
            try:
                return roman_to_int(val)
            except Exception:
                return None
    return None





def center_floating_images(p):
    """
    Tìm tất cả các wp:anchor (ảnh dạng nổi) trong XML của paragraph,
    và cập nhật căn lề ngang của chúng về chính giữa (center).
    """
    try:
        # Tìm tất cả thẻ anchor (local-name là "anchor")
        anchors = p._p.xpath('.//*[local-name()="anchor"]')
        for anchor in anchors:
            # Tìm thẻ positionH (local-name là "positionH")
            posH_list = anchor.xpath('.//*[local-name()="positionH"]')
            for posH in posH_list:
                posH.set('relativeFrom', 'column')
                
                # Xóa posOffset (khoảng cách lệch tĩnh nếu có)
                offsets = posH.xpath('.//*[local-name()="posOffset"]')
                for offset in offsets:
                    posH.remove(offset)
                
                # Cập nhật hoặc chèn align là 'center'
                align_els = posH.xpath('./*[local-name()="align"]')
                if align_els:
                    align_els[0].text = 'center'
                else:
                    # Tạo thẻ align cùng namespace với posH
                    ns_uri = posH.tag.split('}')[0].lstrip('{') if '}' in posH.tag else ''
                    if ns_uri:
                        align = etree.SubElement(posH, '{%s}align' % ns_uri)
                    else:
                        align = etree.SubElement(posH, 'align')
                    align.text = 'center'
    except Exception as e:
        print("Lỗi căn giữa ảnh nổi:", e)

def add_cover_page_border(doc, section):
    """Thêm khung viền trang bìa cho section đầu tiên."""
    try:
        sectPr = section._sectPr
        pgBorders = sectPr.find(qn('w:pgBorders'))
        if pgBorders is not None:
            sectPr.remove(pgBorders)
        
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'text')
        pgBorders.set(qn('w:display'), 'firstPage')
        pgBorders.set(qn('w:zOrder'), 'back') # Không tích "Always display in front"
        
        # Cấu hình khoảng cách viền tính từ văn bản (text): 
        # Top: 20pt, Bottom: 1pt, Left/Right: 4pt
        spacings = {
            'top': '20',
            'bottom': '1',
            'left': '4',
            'right': '4'
        }
        
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12') # size: 1.5 pt
            border.set(qn('w:space'), spacings[side])
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
            
        insert_element_in_order(sectPr, pgBorders, SECTPR_ORDER)

        # Cấu hình settings XML để tắt Surround header / Surround footer / Align borders
        try:
            settings_el = doc.settings.element if hasattr(doc.settings, 'element') else doc.settings._element
            
            # Tắt "Align paragraph borders..." bằng cách xóa alignBordersAndEdges nếu có
            align_el = settings_el.find(qn('w:alignBordersAndEdges'))
            if align_el is not None:
                settings_el.remove(align_el)
                
            # Tắt "Surround header" bằng cách thêm bordersDoNotSurroundHeader
            if settings_el.find(qn('w:bordersDoNotSurroundHeader')) is None:
                no_surround_header = OxmlElement('w:bordersDoNotSurroundHeader')
                insert_element_in_order(settings_el, no_surround_header, SETTINGS_ORDER)
                
            # Tắt "Surround footer" bằng cách thêm bordersDoNotSurroundFooter
            if settings_el.find(qn('w:bordersDoNotSurroundFooter')) is None:
                no_surround_footer = OxmlElement('w:bordersDoNotSurroundFooter')
                insert_element_in_order(settings_el, no_surround_footer, SETTINGS_ORDER)
        except Exception as e:
            print("Lỗi cấu hình surround header/footer:", e)
            
    except Exception as e:
        print("Lỗi khi thêm viền trang bìa:", e)

def get_para_special_type(p):
    """Nhận diện loại tiêu đề đặc biệt của đoạn văn để chia section."""
    style_name = p.style.name if p.style else ''
    if style_name == 'TOC Heading':
        return "toc"
        
    text = p.text.strip()
    if not text:
        return None
    if is_directory_line(text):
        return None
    text_upper = text.upper()
    clean_text = re.sub(r'\s+', ' ', text_upper)
    
    # 1. Lời cảm ơn / Lời cam đoan
    if clean_text in ["LỜI CẢM ƠN", "LỜI CAM ĐOAN", "LOI CAM ON", "LOI CAM DOAN"]:
        return "thanks"
        
    # 2. Mục lục
    if clean_text in ["MỤC LỤC", "MUC LUC"]:
        return "toc"
        
    # 3. Danh mục
    if clean_text in [
        "DANH MỤC HÌNH ẢNH", "DANH MỤC HÌNH", "DANH MỤC BẢNG BIỂU", "DANH MỤC BẢNG",
        "DANH MỤC CÁC THUẬT NGỮ VIẾT TẮT", "DANH MỤC THUẬT NGỮ VIẾT TẮT",
        "DANH MỤC CÁC THUẬT NGỮ, KÝ HIỆU VÀ CÁC CHỮ VIẾT TẮT",
        "DANH MỤC THUẬT NGỮ, KÝ HIỆU VÀ TỪ VIẾT TẮT",
        "DANH MỤC CÁC THUẬT NGỮ, KÝ HIỆU VÀ TỪ VIẾT TẮT",
        "DANH MUC HINH ANH", "DANH MUC HINH", "DANH MUC BANG BIEU", "DANH MUC BANG",
        "DANH MUC CAC THUAT NGU VIET TAT", "DANH MUC THUAT NGU VIET TAT",
        "DANH MUC CAC THUAT NGU, KY HIEU VA CAC CHU VIET TAT",
        "DANH MUC THUAT NGU, KY HIEU VA TU VIET TAT",
        "DANH MUC CAC THUAT NGU, KY HIEU VA TU VIET TAT"
    ]:
        return "list"
        
    # 4. Mở đầu / Chương 1
    if clean_text in ["MỞ ĐẦU", "PHẦN MỞ ĐẦU", "LỜI MỞ ĐẦU", "LỜI NÓI ĐẦU", "MO DAU", "PHAN MO DAU", "LOI MO DAU", "LOI NOI DAU"]:
        return "body"
    if re.match(r'^CHƯƠNG\s+(1|I)\b', clean_text, re.IGNORECASE):
        return "body"
        
    # 5. Tài liệu tham khảo
    if clean_text in ["TÀI LIỆU THAM KHẢO", "TAI LIEU THAM KHAO"]:
        return "references"
        
    return None

def insert_section_break_before(paragraph):
    """Chèn Next Page section break trước paragraph bằng cách tạo paragraph rỗng riêng mang section break.
    Sử dụng paragraph rỗng riêng thay vị paragraph trước để tránh gán section break lên paragraph
    nội dung hoặc tiêu đề danh mục (gây lỗi nội dung TOC bị đẩy sang section sai)."""
    p_element = paragraph._p
    parent = p_element.getparent()
    p_index = parent.index(p_element)
    
    # Kiểm tra xem có bất kỳ section break nào trước đó mà không có văn bản ở giữa không
    for idx in range(p_index - 1, -1, -1):
        sibling = parent[idx]
        if sibling.tag.endswith('}p'):
            existing_pPr = sibling.find(qn('w:pPr'))
            if existing_pPr is not None and existing_pPr.find(qn('w:sectPr')) is not None:
                return False  # Đã có section break gần đó, không chèn thêm
            sibling_text = "".join(sibling.itertext()).strip()
            if sibling_text:
                break
        elif sibling.tag.endswith('}sdt'):
            sdt_text = "".join(sibling.itertext()).strip()
            if sdt_text:
                break
    
    # Tạo paragraph rỗng riêng để mang section break
    new_p = OxmlElement('w:p')
    parent.insert(parent.index(p_element), new_p)
    
    pPr = new_p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    insert_element_in_order(pPr, sectPr, PPR_ORDER)
    return True

def partition_paragraphs_by_section(doc):
    """Phân nhóm paragraphs theo từng section 1-to-1."""
    sections_paragraphs = []
    current_list = []
    for p in doc.paragraphs:
        current_list.append(p)
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:sectPr')) is not None:
            sections_paragraphs.append(current_list)
            current_list = []
    sections_paragraphs.append(current_list)
    return sections_paragraphs

def classify_sections(doc):
    """Gán nhãn loại section cho từng section trong tài liệu."""
    parts = partition_paragraphs_by_section(doc)
    section_types = []
    current_type = "cover"
    for part in parts:
        part_type = None
        for p in part:
            p_type = get_para_special_type(p)
            if p_type is not None:
                part_type = p_type
                break
        if part_type is not None:
            current_type = part_type
        section_types.append(current_type)
    return section_types

def get_element_section_types(doc):
    """
    Trích xuất loại section cho từng element (paragraph hoặc table) trong body.
    """
    body_elements = doc.element.body
    
    element_section_indices = {}
    current_section_idx = 0
    
    for child in body_elements:
        if child.tag.endswith('}p'):
            has_sect = False
            pPr = next((c for c in child if c.tag.endswith('}pPr')), None)
            if pPr is not None:
                sectPr = next((c for c in pPr if c.tag.endswith('}sectPr')), None)
                if sectPr is not None:
                    has_sect = True
            
            element_section_indices[child] = current_section_idx
            if has_sect:
                current_section_idx += 1
        elif child.tag.endswith('}tbl'):
            element_section_indices[child] = current_section_idx
            
    return element_section_indices

def set_section_page_numbering(section, fmt=None, start=None):
    """Cấu hình định dạng và số trang bắt đầu cho section trong w:sectPr."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        insert_element_in_order(sectPr, pgNumType, SECTPR_ORDER)
    
    if fmt is not None:
        pgNumType.set(qn('w:fmt'), fmt)
    else:
        if qn('w:fmt') in pgNumType.attrib:
            del pgNumType.attrib[qn('w:fmt')]
            
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    else:
        if qn('w:start') in pgNumType.attrib:
            del pgNumType.attrib[qn('w:start')]



def parent_already_has_toc(parent, p_index, indicator):
    """Kiểm tra xem trong vòng 2 phần tử tiếp theo đã có trường TOC tương ứng chưa."""
    from lxml import etree
    xpath_func = etree.XPath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    for i in range(1, 3):
        if p_index + i < len(parent):
            el = parent[p_index + i]
            instrs = xpath_func(el)
            for instr in instrs:
                if instr.text and 'TOC' in instr.text and indicator in instr.text:
                    return True
    return False


def insert_toc_field_after(p, doc, field_instr, style_name='Table of Figures', font_name='Times New Roman'):
    """Chèn mã trường TOC tự động của Word vào ngay sau paragraph tiêu đề danh mục."""
    try:
        from docx.text.paragraph import Paragraph
        parent = p._p.getparent()
        new_p_el = OxmlElement('w:p')
        
        # Chèn paragraph mới vào dưới đoạn tiêu đề trong XML trước
        p_index = parent.index(p._p)
        parent.insert(p_index + 1, new_p_el)
        
        # Tạo wrapper Paragraph để cấu hình an toàn các paragraph format bằng python-docx API
        # nhằm tránh vi phạm thứ tự XML Schema (Word XML Sequence) gây lỗi file corrupt.
        p_wrapped = Paragraph(new_p_el, doc)
        
        # Áp dụng style nếu tồn tại trong tài liệu
        if style_name in doc.styles:
            try:
                p_wrapped.style = doc.styles[style_name]
            except Exception as e:
                    logging.warning(f"Silent error ignored: {e}")
                
        # Căn đều 2 bên (Justify)
        p_wrapped.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Giãn dòng 1.5 và cách đoạn 0pt
        p_wrapped.paragraph_format.line_spacing = 1.5
        p_wrapped.paragraph_format.space_before = Pt(0)
        p_wrapped.paragraph_format.space_after = Pt(0)
        
        # Không thụt lề
        p_wrapped.paragraph_format.first_line_indent = Pt(0)
        p_wrapped.paragraph_format.left_indent = Pt(0)
        p_wrapped.paragraph_format.right_indent = Pt(0)
        
        # Kích hoạt không cách giữa các đoạn cùng style
        try:
            set_contextual_spacing(p_wrapped, True)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
                
        # Thêm run chứa mã trường TOC và gán các character properties an toàn bằng python-docx API
        run = p_wrapped.add_run()
        set_run_font(run, font_name, 14, bold=False, italic=False)
        
        # Thiết lập rõ ràng w:val="0" cho các thuộc tính bold/italic trong XML để ghi đè mọi kế thừa từ mẫu Word gốc
        rPr = run.element.get_or_add_rPr()
        for tag in ['w:b', 'w:bCs', 'w:i', 'w:iCs']:
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn('w:val'), '0')
        
        # Xóa thẻ w:t mặc định mà add_run tạo ra
        t_el = run._r.find(qn('w:t'))
        if t_el is not None:
            run._r.remove(t_el)
            
        # Thêm các thẻ XML trường động vào run
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_instr
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    except Exception as e:
        print("Lỗi chèn TOC tự động:", e)


def make_caption_paragraph(p, label, chapter, seq_num, text_content, font_name, font_size):
    """
    Tái tạo paragraph chú thích sử dụng trường SEQ động của Word để tương thích
    hoàn toàn với danh mục tự động (References -> Insert Caption -> Label).
    Định dạng: [label] [chapter].[SEQ].[text_content]
    Ví dụ: Hình 2.1. Tên hình
    """
    p.text = ""  # Xóa các runs cũ
    
    # 1. Phần nhãn và chương (Ví dụ: "Hình 2.")
    r_prefix = p.add_run(f"{label} {chapter}.")
    
    # 2. Trường SEQ của Word (Chèn mã trường SEQ động vào một run duy nhất để tránh lỗi XML của Word)
    r_field = p.add_run()
    # Xóa thẻ w:t trống mà add_run tạo ra mặc định
    t_el = r_field._r.find(qn('w:t'))
    if t_el is not None:
        r_field._r.remove(t_el)
        
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f" SEQ {label} \\* ARABIC \\s 1 "
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    t_val = OxmlElement('w:t')
    t_val.text = str(seq_num)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_field._r.append(fldChar1)
    r_field._r.append(instrText)
    r_field._r.append(fldChar2)
    r_field._r.append(t_val)
    r_field._r.append(fldChar3)
    
    # 3. Phần mô tả phía sau (Ví dụ: ". Logo công ty")
    r_text = p.add_run(f". {text_content}")
    
    # Áp dụng font cỡ chữ 14, nghiêng, không in đậm cho chú thích hình/bảng
    for r in p.runs:
        set_run_font(r, font_name, 14, bold=False, italic=True)
        rPr = r.element.get_or_add_rPr()
        for tag in ['w:b', 'w:bCs']:
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn('w:val'), '0')
        for tag in ['w:i', 'w:iCs']:
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn('w:val'), '1')


def is_directory_line(text):
    """
    Kiểm tra xem một dòng có phải là một mục của mục lục hoặc danh mục hay không
    (dựa trên sự xuất hiện của dấu chấm lửng ... hoặc số trang ở cuối dòng kèm tab/nhiều khoảng trắng).
    """
    # Nếu chứa dấu chấm lửng (từ 2 dấu chấm liên tiếp trở lên)
    if '..' in text:
        return True
    # Nếu kết thúc bằng Tab và số trang
    if re.search(r'\t\s*\d+$', text):
        return True
    # Nếu kết thúc bằng ít nhất 2 khoảng trắng và số trang
    if re.search(r'\s{2,}\d+$', text):
        return True
    if re.search(r'\.+\s*\d+$', text):
        return True
    return False


def is_toc_or_directory_paragraph(p):
    """
    Kiểm tra xem một paragraph có thuộc mục lục hoặc danh mục hay không
    dựa trên style name, các liên kết (hyperlinks) _Toc hoặc định dạng nội dung.
    """
    text = p.text.strip()
    if not text:
        return False
    
    # 1. Kiểm tra tên style của paragraph
    style_name = p.style.name if p.style else ''
    style_name_upper = style_name.upper()
    if style_name_upper.startswith('TOC') or 'TABLE OF FIGURES' in style_name_upper or 'TABLEOFFIGURES' in style_name_upper:
        return True
    
    # 2. Kiểm tra các hyperlink trong XML (thường là trường TOC có thuộc tính anchor bắt đầu bằng _Toc)
    p_el = p._p
    hyperlinks = p_el.xpath('.//*[local-name()="hyperlink"]')
    for hl in hyperlinks:
        anchor = hl.get(qn('w:anchor'))
        if anchor and anchor.startswith('_Toc'):
            return True

    # 3. Kiểm tra định dạng dòng của mục lục/danh mục
    if is_directory_line(text):
        return True
        
    return False


def remove_manual_entries_under_titles(doc):
    """
    Xóa các dòng chú thích thủ công (không phải trường TOC động) nằm ngay dưới
    các tiêu đề DANH MỤC HÌNH ẢNH hoặc DANH MỤC BẢNG BIỂU để tránh trùng lặp.
    """
    paragraphs_to_remove = []
    
    i = 0
    paragraphs = list(doc.paragraphs)
    n = len(paragraphs)
    while i < n:
        p = paragraphs[i]
        text_upper = p.text.strip().upper()
        clean_text = re.sub(r'\s+', ' ', text_upper)
        
        # Nếu gặp tiêu đề danh mục hình ảnh, bảng biểu hoặc mục lục
        if clean_text in ["DANH MỤC HÌNH ẢNH", "DANH MỤC HÌNH", "DANH MỤC BẢNG BIỂU", "DANH MỤC BẢNG", "MỤC LỤC"]:
            j = i + 1
            while j < n:
                next_p = paragraphs[j]
                next_text = next_p.text.strip()
                
                if not next_text:
                    j += 1
                    continue
                    
                next_text_upper = next_text.upper()
                next_style = next_p.style.name if next_p.style else 'Normal'
                
                is_end_of_directory = False
                if next_style == 'Heading 1' and next_text_upper not in ["DANH MỤC HÌNH ẢNH", "DANH MỤC HÌNH", "DANH MỤC BẢNG BIỂU", "DANH MỤC BẢNG", "MỤC LỤC"]:
                    is_end_of_directory = True
                elif (re.match(r'^CHƯƠNG\s+[IVX\d]+', next_text, re.IGNORECASE) or next_text_upper in ["MỞ ĐẦU", "PHẦN MỞ ĐẦU", "LỜI MỞ ĐẦU", "LỜI NÓI ĐẦU"]) and not is_directory_line(next_text):
                    is_end_of_directory = True
                elif next_style in ['Heading 2', 'Heading 3'] and not is_directory_line(next_text):
                    is_end_of_directory = True
                    
                if is_end_of_directory:
                    break
                    
                # Kiểm tra xem có chứa trường TOC động không
                has_toc_field = False
                from lxml import etree
                xpath_func = etree.XPath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                instrs = xpath_func(next_p._p)
                for instr in instrs:
                    if instr.text and 'TOC' in instr.text:
                        has_toc_field = True
                        break
                        
                if not has_toc_field:
                    if is_directory_line(next_text) or re.match(r'^(Hình|Ảnh|Bảng|Chương|CHƯƠNG|\d+(\.\d+)*)\b', next_text, re.IGNORECASE):
                        paragraphs_to_remove.append(next_p)
                
                j += 1
            i = j
        else:
            i += 1
            
    for p in paragraphs_to_remove:
        try:
            p_el = p._p
            p_el.getparent().remove(p_el)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")


def clean_paragraph_whitespace(p):
    """
    Dọn dẹp khoảng trắng thừa trong paragraph mà không làm hỏng định dạng các run (in đậm, in nghiêng, v.v.).
    """
    if not p.runs:
        return
        
    # 1. Xóa khoảng trắng/tab ở đầu paragraph (ở run đầu tiên có chữ)
    for r in p.runs:
        if r.text:
            stripped = r.text.lstrip(' \t\n\r ')
            if stripped != r.text:
                r.text = stripped
            if r.text:
                break
                
    # 2. Xóa khoảng trắng/tab ở cuối paragraph (ở run cuối cùng có chữ)
    for r in reversed(p.runs):
        if r.text:
            stripped = r.text.rstrip(' \t\n\r ')
            if stripped != r.text:
                r.text = stripped
            if r.text:
                break
                
    # 3. Thu gọn nhiều khoảng trắng liên tiếp bên trong mỗi run
    for r in p.runs:
        if r.text:
            cleaned = re.sub(r' {2,}', ' ', r.text)
            cleaned = re.sub(r'\xa0{2,}', ' ', cleaned)
            if cleaned != r.text:
                r.text = cleaned


def cleanup_document_whitespace(doc, skip_paras=None):
    """
    Bước tiền xử lý: Dọn dẹp khoảng trắng thừa và xóa các paragraph rỗng không chứa ảnh hoặc section break.
    Không xóa các đoạn trống thuộc trang bìa (cover) để tránh hỏng bố cục.
    """
    if skip_paras is None:
        skip_paras = set()
    try:
        parts = partition_paragraphs_by_section(doc)
        section_types = classify_sections(doc)
        
        cover_paragraphs = set()
        for idx, part in enumerate(parts):
            sect_type = section_types[idx] if idx < len(section_types) else "body"
            if sect_type == "cover":
                cover_paragraphs.update(part)
    except Exception as e:
        print("Lỗi khi phân loại section để dọn dẹp dòng trống:", e)
        cover_paragraphs = set()

    for p in doc.paragraphs:
        if p._p in skip_paras:
            continue
        if is_directory_line(p.text):
            continue
        clean_paragraph_whitespace(p)

    paras_to_remove = []
    cover_elements = {x._p for x in cover_paragraphs}
    for p in doc.paragraphs:
        if p._p in skip_paras:
            continue
        if p._p in cover_elements:
            continue
            
        if not p.text.strip() and not paragraph_has_image(p):
            pPr = p._p.find(qn('w:pPr'))
            has_sect = False
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                has_sect = True
            
            if not has_sect:
                paras_to_remove.append(p)
                
    for p in paras_to_remove:
        try:
            p_el = p._p
            parent = p_el.getparent()
            if parent is not None:
                parent.remove(p_el)
        except Exception as e:
            print("Lỗi khi xóa đoạn văn rỗng:", e)


def get_heading_level_from_style(style_val):
    if not style_val:
        return None
    m = re.search(r'Heading\s*(\d)', style_val, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return None


def split_paragraphs_on_soft_breaks(doc, skip_paras=None):
    """
    Tìm các paragraph có chứa ký tự xuống dòng mềm (\n) và tách chúng thành các đoạn văn riêng biệt.
    Đồng thời tự động nâng cấp style của phần tiêu đề nếu có chứa thuộc tính Heading trong các run.
    """
    if skip_paras is None:
        skip_paras = set()
    import copy
    paragraphs = list(doc.paragraphs)
    
    for i in range(len(paragraphs) - 1, -1, -1):
        p = paragraphs[i]
        if p._p in skip_paras:
            continue
        text = p.text
        if "\n" not in text:
            continue
            
        p_el = p._p
        parent = p_el.getparent()
        p_parent = p._parent
        orig_style_name = p.style.name
        
        # Nhóm các run theo từng dòng văn bản
        lines_runs = []
        current_line = []
        
        for r in p.runs:
            r_text = r.text
            if "\n" in r_text:
                parts = r_text.split("\n")
                for idx, part in enumerate(parts):
                    if idx > 0:
                        lines_runs.append(current_line)
                        current_line = []
                    if part:
                        new_r_el = copy.deepcopy(r._r)
                        w_t = new_r_el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        if w_t is not None:
                            w_t.text = part
                        from docx.text.run import Run
                        new_run = Run(new_r_el, p)
                        current_line.append(new_run)
            else:
                current_line.append(r)
                
        if current_line:
            lines_runs.append(current_line)
            
        if len(lines_runs) <= 1:
            continue
            
        # Xóa các runs cũ trong paragraph ban đầu và cập nhật dòng đầu tiên
        p_el_clear = p._p
        for r_el in list(p_el_clear.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')):
            p_el_clear.remove(r_el)
            
        first_line = lines_runs[0]
        for r in first_line:
            p_el_clear.append(r._r)
            
        # Nhận diện nếu dòng đầu tiên có style run chứa Heading
        first_line_heading_level = None
        for r in first_line:
            rPr = r._r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                rStyle = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rStyle')
                if rStyle is not None:
                    style_val = rStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    lvl = get_heading_level_from_style(style_val)
                    if lvl is not None:
                        first_line_heading_level = lvl
                        break
                        
        # Gán style cho đoạn đầu tiên
        if orig_style_name.startswith('Heading'):
            pass
        elif first_line_heading_level is not None:
            target_style = f"Heading {first_line_heading_level}"
            try:
                p.style = target_style
            except Exception:
                pPr = p_el_clear.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                if pPr is not None:
                    pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                    if pStyle is not None:
                        pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', target_style.replace(" ", ""))
                        
        # Thêm các đoạn văn phía sau
        current_p_el = p_el
        for line_idx in range(1, len(lines_runs)):
            runs_in_line = lines_runs[line_idx]
            new_p_el = OxmlElement('w:p')
            
            pPr = p_el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                new_pPr = copy.deepcopy(pPr)
                # Xóa style cũ của đoạn văn gốc để gán style mới
                pStyle = new_pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if pStyle is not None:
                    new_pPr.remove(pStyle)
                new_p_el.append(new_pPr)
                
            for r in runs_in_line:
                new_p_el.append(r._r)
                
            current_p_el.addnext(new_p_el)
            current_p_el = new_p_el
            
            new_p = Paragraph(new_p_el, p_parent)
            # Dòng tiếp theo không phải là heading mà là body text (Normal hoặc style gốc nếu gốc không phải heading)
            target_style_name = 'Normal' if orig_style_name.startswith('Heading') else orig_style_name
            try:
                new_p.style = target_style_name
            except Exception:
                new_pPr = new_p_el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                if new_pPr is None:
                    new_pPr = OxmlElement('w:pPr')
                    new_p_el.insert(0, new_pPr)
                pStyle = new_pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    new_pPr.append(pStyle)
                pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', target_style_name.replace(" ", ""))


def merge_split_chapter_titles(doc, skip_paras=None):
    """
    Tìm các chương bị tách thành 2 dòng (Dòng 1: CHƯƠNG X, Dòng 2: Tên chương)
    và gộp chúng thành 1 paragraph chứa ký tự xuống dòng mềm (\n).
    """
    if skip_paras is None:
        skip_paras = set()
    paragraphs = list(doc.paragraphs)
    i = 0
    n = len(paragraphs)
    paras_to_remove = []
    
    while i < n - 1:
        p1 = paragraphs[i]
        if p1._p in skip_paras:
            i += 1
            continue
        text1 = p1.text.strip()
        
        # Nhận diện dòng 1 chỉ chứa "CHƯƠNG X" hoặc "CHƯƠNG X."
        if re.match(r'^CHƯƠNG\s+[IVX\d]+[\.:]?$', text1, re.IGNORECASE):
            # Tìm dòng tiếp theo không rỗng để làm tên chương
            j = i + 1
            p2 = None
            while j < n:
                candidate = paragraphs[j]
                if candidate.text.strip():
                    p2 = candidate
                    break
                j += 1
                
            if p2 is not None:
                text2 = p2.text.strip()
                # Tên chương không quá dài, không phải là heading khác
                if len(text2) < 150 and not re.match(r'^(CHƯƠNG|Heading|MỤC LỤC|DANH MỤC)\b', text2, re.IGNORECASE):
                    # Tiến hành gộp trực tiếp văn bản bằng cách chèn \n
                    p1.text = f"{text1}\n{text2}"
                    paras_to_remove.append(p2)
                    i = j
                    continue
        i += 1
        
    for p in paras_to_remove:
        try:
            p_el = p._p
            parent = p_el.getparent()
            if parent is not None:
                parent.remove(p_el)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")


def get_cover_table(doc):
    """
    Tìm và trả về table trang bìa nếu có.
    Table trang bìa thường là table đầu tiên của tài liệu, nằm ở section đầu tiên
    và chứa các từ khóa đặc trưng của trang bìa.
    """
    if not doc.tables:
        return None
    
    first_table = doc.tables[0]
    
    # Kiểm tra xem table này có chứa các từ khóa đặc trưng của bìa không
    text_content = ""
    for row in first_table.rows:
        for cell in row.cells:
            text_content += cell.text + " "
            
    text_upper = text_content.upper()
    keywords = ["TRƯỜNG ĐẠI HỌC", "BỘ CÔNG THƯƠNG", "BỘ GIÁO DỤC", "ĐỒ ÁN", "KHOÁ LUẬN", "LUẬN VĂN", "BÁO CÁO", "GVHD", "GVPB", "SINH VIÊN", "HÀ NỘI"]
    match_count = sum(1 for kw in keywords if kw in text_upper)
    
    # Nếu chứa từ 2 từ khóa trở lên, khả năng cao là table trang bìa
    if match_count >= 2:
        return first_table
    return None


def format_document(input_path, output_path, opts):
    """
    Định dạng tài liệu Word (.docx) dựa trên dictionary cấu hình opts.
    Hỗ trợ: căn lề trang, phông chữ, dãn dòng, tiêu đề, bảng biểu,
    các thành phần hành chính theo Nghị định 30/2020/NĐ-CP, đánh số trang.
    """
    doc = Document(input_path)
    cover_table = get_cover_table(doc)
    format_cover = opts.get('format_cover', True)

    # --- PHÂN LOẠI SECTION TRANG BÌA ĐỂ THIẾT LẬP BỎ QUA ---
    cover_paras = []
    cover_section_indices = set()
    section_types = []
    try:
        parts = partition_paragraphs_by_section(doc)
        section_types = classify_sections(doc)
        for idx, part in enumerate(parts):
            sect_type = section_types[idx] if idx < len(section_types) else "body"
            if sect_type == "cover":
                cover_paras.extend(part)
                cover_section_indices.add(idx)
    except Exception as e:
        print("Lỗi khi phân tách section trang bìa:", e)
        cover_paras = []
        cover_section_indices = set()

    skip_paras = {p._p for p in cover_paras} if not format_cover else set()

    # --- 0. XÓA SẠCH MỌI SỐ TRANG CŨ VÀ HEADERS/FOOTERS ---
    try:
        clear_all_headers_footers(doc, skip_sections=cover_section_indices)
        try:
            settings_el = doc.settings.element if hasattr(doc.settings, 'element') else doc.settings._element
            even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
            if even_odd is not None:
                settings_el.remove(even_odd)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")
    except Exception as e:
        print("Lỗi khi xóa sạch header/footer cũ:", e)

    # --- 0. TÁCH CÁC DÒNG XUỐNG DÒNG MỀM CHỨA HEADING / PHÂN TÁCH PARAGRAPH ---
    try:
        split_paragraphs_on_soft_breaks(doc, skip_paras=skip_paras)
    except Exception as e:
        print("Lỗi khi tách soft break:", e)

    # --- 0. GỘP TIÊU ĐỀ CHƯƠNG BỊ TÁCH DÒNG ---
    try:
        merge_split_chapter_titles(doc, skip_paras=skip_paras)
    except Exception as e:
        print("Lỗi khi gộp tiêu đề chương:", e)

    # --- 0. DỌN DẸP KHOẢNG TRẤNG VÀ PARAGRAPH RỖNG TRƯỚC ---
    try:
        cleanup_document_whitespace(doc, skip_paras=skip_paras)
    except Exception as e:
        print("Lỗi khi dọn dẹp khoảng trắng:", e)

    # --- 0. DỌN DẸP DÒNG CHÚ THÍCH THỦ CÔNG DƯỚI TIÊU ĐỀ DANH MỤC TRƯỚC ---
    try:
        remove_manual_entries_under_titles(doc)
    except Exception as e:
        print("Lỗi khi dọn dẹp danh mục thủ công:", e)

    # --- 0. PHÂN CHIA SECTION BREAKS DỰA TRÊN CẤU TRÚC VĂN BẢN ---
    if opts.get('add_page_numbers', True):
        # Dọn dẹp tất cả các section break cũ trong paragraph để tránh trùng lặp khi định dạng lại nhiều lần
        for p in doc.paragraphs:
            if p._p in skip_paras:
                continue
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    pPr.remove(sectPr)

        trigger_paras = []
        last_type = "cover"
        ORDER = ["cover", "thanks", "toc", "list", "body", "references"]
        for p in doc.paragraphs:
            p_type = get_para_special_type(p)
            if p_type is not None:
                if p_type in ORDER and ORDER.index(p_type) > ORDER.index(last_type):
                    trigger_paras.append(p)
                    last_type = p_type
        for p in trigger_paras:
            insert_section_break_before(p)

    # --- 1. THIẾT LẬP TRANG (Margins & Khổ A4) ---
    margin_top = float(opts.get('margin_top', 20))
    margin_bottom = float(opts.get('margin_bottom', 20))
    margin_left = float(opts.get('margin_left', 30))
    margin_right = float(opts.get('margin_right', 15))

    for idx, section in enumerate(doc.sections):
        if not format_cover and idx in cover_section_indices:
            continue
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(margin_top)
        section.bottom_margin = Mm(margin_bottom)
        section.left_margin = Mm(margin_left)
        section.right_margin = Mm(margin_right)
        # Căn giữa hoàn hảo số trang Header/Footer theo chiều dọc (trừ đi chiều cao ước lượng của dòng chữ là 5mm)
        section.header_distance = Mm(max(1.0, (margin_top - 5) / 2))
        section.footer_distance = Mm(max(1.0, (margin_bottom - 5) / 2))

    # --- 1.5. ĐỊNH DẠNG TRANG BÌA (COVER PAGE) ---
    if format_cover and cover_paras and cover_table is None:
        try:
            # Thêm viền trang bìa cho section đầu tiên
            add_cover_page_border(doc, doc.sections[0])
            
            # 1. Tìm các paragraph chứa ảnh (logo) và tách chúng ra
            image_paras = [p for p in cover_paras if paragraph_has_image(p)]
            for p in image_paras:
                p_el = p._p
                parent = p_el.getparent()
                if parent is not None:
                    parent.remove(p_el)
                    
            # Loại bỏ image_paras khỏi cover_paras để xử lý text
            cover_paras = [p for p in cover_paras if p not in image_paras]
            
            # Trích xuất văn bản và cấu trúc khối từ cover_paras
            separator_idx = -1
            for idx, p in enumerate(cover_paras):
                t = p.text.strip()
                # Hỗ trợ ======***====== và các loại đường phân cách khác
                if re.match(r'^[\-\_\=\.\s\*~•]{5,}$', t):
                    separator_idx = idx
                    break

            info_indices = []
            for idx, p in enumerate(cover_paras):
                t = p.text.strip().upper()
                if any(kw in t for kw in ["GVHD", "GVPB", "SINH VIÊN", "MÃ SINH VIÊN", "MÃ SỐ SINH VIÊN", "LỚP", "KHÓA", "MSSV", "MSV", "GIÁO VIÊN HƯỚNG DẪN"]):
                    info_indices.append(idx)

            location_idx = -1
            for idx in range(len(cover_paras) - 1, -1, -1):
                if cover_paras[idx].text.strip():
                    if idx not in info_indices:
                        location_idx = idx
                        break

            school_header_paras = []
            if separator_idx != -1:
                school_header_paras = cover_paras[:separator_idx + 1]
            else:
                for idx, p in enumerate(cover_paras):
                    t = p.text.strip()
                    if not t:
                        continue
                    t_upper = t.upper()
                    if any(kw in t_upper for kw in ["BỘ", "TRƯỜNG", "ĐẠI HỌC", "VIỆN", "KHOA"]):
                        school_header_paras.append(p)
                    else:
                        break

            info_paras = [cover_paras[i] for i in info_indices]
            location_para = cover_paras[location_idx] if location_idx != -1 else None

            start_middle = separator_idx + 1 if separator_idx != -1 else len(school_header_paras)
            end_middle = info_indices[0] if info_indices else len(cover_paras)
            middle_paras_raw = cover_paras[start_middle:end_middle]

            def strip_paragraph_list(para_list):
                start = 0
                while start < len(para_list) and not para_list[start].text.strip():
                    start += 1
                end = len(para_list)
                while end > start and not para_list[end - 1].text.strip():
                    end -= 1
                return para_list[start:end]

            middle_paras = strip_paragraph_list(middle_paras_raw)

            # Xây dựng sequence mục tiêu dựa trên các dữ liệu trích xuất được
            sequence = []
            
            # A. School Header
            for p in school_header_paras:
                t = p.text.strip()
                t_upper = t.upper()
                is_separator = re.match(r'^[\-\_\=\.\s\*~•]{5,}$', t)
                if is_separator:
                    sequence.append({'text': t, 'type': 'school', 'bold': False})
                elif "BỘ" in t_upper and "TRƯỜNG" not in t_upper and "ĐẠI HỌC" not in t_upper:
                    sequence.append({'text': t, 'type': 'school', 'bold': False})
                else:
                    sequence.append({'text': t, 'type': 'school', 'bold': True})
            
            # B. Gap 1: 5 empty lines (hoặc chứa logo trường ở giữa)
            if image_paras:
                sequence.append({'text': '', 'type': 'empty'})
                sequence.append({'text': '', 'type': 'empty'})
                for logo_p in image_paras:
                    sequence.append({'type': 'image', 'para_obj': logo_p})
                sequence.append({'text': '', 'type': 'empty'})
                sequence.append({'text': '', 'type': 'empty'})
            else:
                for _ in range(5):
                    sequence.append({'text': '', 'type': 'empty'})
                
            # C. Middle block (Giữ nguyên thứ tự các dòng Đề tài, Ngành, Loại đồ án và dòng trống nguyên bản)
            for p in middle_paras:
                t = p.text.strip()
                if not t:
                    sequence.append({'text': '', 'type': 'empty'})
                else:
                    sequence.append({'text': t, 'type': 'title', 'bold': True})
                
            # D. Gap 3: 3 empty lines
            for _ in range(3):
                sequence.append({'text': '', 'type': 'empty'})
                
            # E. Info Block
            for p in info_paras:
                sequence.append({'text': p.text.strip(), 'type': 'info'})
                
            # F. Gap 4: dòng trống động để đẩy địa danh xuống dòng cuối cùng của trang đầu (dòng thứ 29)
            # Ước lượng số dòng thực tế để tránh nhảy sang trang 2 nếu tiêu đề dài hoặc có ảnh logo
            title_extra_lines = 0
            for p_mid in middle_paras:
                t_len = len(p_mid.text.strip())
                if t_len > 40:
                    title_extra_lines += (t_len // 40)
                    
            logo_offset = 5 if image_paras else 0
            # Add +1 to used_lines for the dedicated empty paragraph at the very end of Section 0 (to hold the section break)
            used_lines = len(school_header_paras) + 5 + len(middle_paras) + 3 + len(info_paras) + 1 + logo_offset + title_extra_lines + 1
            gap4_lines = max(1, 29 - used_lines)
            for _ in range(gap4_lines):
                sequence.append({'text': '', 'type': 'empty'})
                
            # G. Location / Year
            if location_para is not None:
                loc_clean = re.sub(r'\s+', ' ', location_para.text).strip()
                sequence.append({'text': loc_clean, 'type': 'location', 'bold': False})
                
            # H. Dedicated empty paragraph at the end to hold the sectPr (section break)
            sequence.append({'text': '', 'type': 'empty'})

            N = len(sequence)
            
            # Điều chỉnh số lượng paragraphs trong cover_paras để khớp với sequence target
            last_p = cover_paras[-1]
            while len(cover_paras) < N:
                new_p = last_p.insert_paragraph_before("")
                cover_paras.insert(-1, new_p)
                
            while len(cover_paras) > N:
                p_to_remove = cover_paras[len(cover_paras) - 2]
                p_el = p_to_remove._p
                p_el.getparent().remove(p_el)
                cover_paras.pop(len(cover_paras) - 2)

            # Điền dữ liệu và áp dụng định dạng Times New Roman, 14pt, giãn dòng 1.5, 0pt trước/sau
            font_name_cover = opts.get('font_name', 'Times New Roman')
            for idx in range(N):
                item = sequence[idx]
                
                if item['type'] == 'image':
                    # Phục hồi ảnh logo vào đúng vị trí XML và cập nhật cover_paras
                    logo_p = item['para_obj']
                    dummy_p = cover_paras[idx]
                    dummy_el = dummy_p._p
                    parent = dummy_el.getparent()
                    pos = parent.index(dummy_el)
                    
                    parent.insert(pos, logo_p._p)
                    parent.remove(dummy_el)
                    cover_paras[idx] = logo_p
                    
                    # Căn giữa logo và định dạng spacing
                    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_p.paragraph_format.first_line_indent = Pt(0)
                    logo_p.paragraph_format.left_indent = Pt(0)
                    logo_p.paragraph_format.right_indent = Pt(0)
                    logo_p.paragraph_format.space_before = Pt(0)
                    logo_p.paragraph_format.space_after = Pt(0)
                    logo_p.paragraph_format.line_spacing = 1.5
                    
                    # Căn giữa tất cả ảnh nổi của logo
                    center_floating_images(logo_p)
                    continue
                    
                # Với các paragraph thông thường
                p = cover_paras[idx]
                p.text = ""  # Xóa các runs cũ
                
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.right_indent = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                
                text = item['text']
                p_type = item['type']
                
                if p_type == 'empty':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.left_indent = Pt(0)
                    r = p.add_run()
                    set_run_font(r, font_name_cover, 14, bold=False, italic=False)
                    continue
                    
                if p_type == 'school':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.left_indent = Pt(0)
                    t_val = text.upper() if not re.match(r'^[\-\_\=\.\s\*~•]{5,}$', text) else text
                    r = p.add_run(t_val)
                    set_run_font(r, font_name_cover, 14, bold=item['bold'], italic=False)
                    
                elif p_type == 'title':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.left_indent = Pt(0)
                    r = p.add_run(text.upper())
                    set_run_font(r, font_name_cover, 14, bold=item['bold'], italic=False)
                    
                elif p_type == 'info':
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Mm(47.5)
                    
                    # Cấu hình các tab stops: 3.5 cm cho cột giá trị 1, 8.0 cm cho cột nhãn 2, 9.5 cm cho cột giá trị 2
                    try:
                        p.paragraph_format.tab_stops.clear_all()
                        p.paragraph_format.tab_stops.add_tab_stop(Mm(35))
                        p.paragraph_format.tab_stops.add_tab_stop(Mm(80))
                        p.paragraph_format.tab_stops.add_tab_stop(Mm(95))
                    except Exception as e:
                            logging.warning(f"Silent error ignored: {e}")
                        
                    # Chuẩn hóa khoảng trắng và tab sau dấu hai chấm để tránh chia nhỏ thành nhiều tab lệch nhau
                    normalized_text = re.sub(r':\s*[\t\s]+', ':\t', text)
                    last_label = None
                    parts = re.split(r'(\t|\s{2,})', normalized_text)
                    for part in parts:
                        if not part:
                            continue
                        if re.match(r'^(\t|\s{2,})$', part):
                            if last_label is not None:
                                is_mssv = any(k in last_label.lower() for k in ["mã số sinh viên", "mã sinh viên", "mssv", "msv", "mã sv"])
                                tab_str = "\t" if is_mssv else ("\t\t" if len(last_label) < 12 else "\t")
                                r = p.add_run(tab_str)
                                last_label = None
                            else:
                                r = p.add_run("\t")
                            set_run_font(r, font_name_cover, 14, bold=False, italic=False)
                        elif ':' in part:
                            label, val = part.split(':', 1)
                            label_clean = label.strip()
                            val_clean = val.strip()
                            
                            r_lbl = p.add_run(label_clean + ":")
                            set_run_font(r_lbl, font_name_cover, 14, bold=True, italic=False)
                            
                            if val_clean:
                                is_mssv = any(k in label_clean.lower() for k in ["mã số sinh viên", "mã sinh viên", "mssv", "msv", "mã sv"])
                                tab_str = "\t" if is_mssv else ("\t\t" if len(label_clean) < 12 else "\t")
                                r_tab = p.add_run(tab_str)
                                set_run_font(r_tab, font_name_cover, 14, bold=False, italic=False)
                                r_val = p.add_run(val_clean)
                                set_run_font(r_val, font_name_cover, 14, bold=False, italic=False)
                                last_label = None
                            else:
                                last_label = label_clean
                        else:
                            r = p.add_run(part.strip())
                            set_run_font(r, font_name_cover, 14, bold=False, italic=False)
                            
                elif p_type == 'location':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.left_indent = Pt(0)
                    loc_formatted = text.replace('–', '-').title()
                    loc_formatted = re.sub(r'\s+', ' ', loc_formatted).strip()
                    r = p.add_run(loc_formatted)
                    set_run_font(r, font_name_cover, 14, bold=item['bold'], italic=False)
        except Exception as e:
            print("Lỗi khi định dạng trang bìa:", e)

    # --- 2. ĐỌC CẤU HÌNH ---
    font_name = opts.get('font_name', 'Times New Roman')
    body_size = float(opts.get('body_size', 13))
    line_spacing = float(opts.get('line_spacing', 1.3))
    space_before = float(opts.get('space_before', 0))
    space_after = float(opts.get('space_after', 6))
    indent_val = float(opts.get('first_line_indent', 10))
    first_line_indent = Mm(indent_val) if indent_val > 0 else None
    contextual_spacing = opts.get('contextual_spacing', True)

    align_str = opts.get('alignment', 'justify')
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    body_alignment = alignment_map.get(align_str, WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Cập nhật style Normal
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = font_name
        normal_style.font.size = Pt(body_size)
        normal_style.paragraph_format.line_spacing = line_spacing
        normal_style.paragraph_format.space_before = Pt(space_before)
        normal_style.paragraph_format.space_after = Pt(space_after)
        normal_style.paragraph_format.first_line_indent = first_line_indent
        set_contextual_spacing(normal_style, contextual_spacing)
    except Exception as e:
            logging.warning(f"Silent error ignored: {e}")

    # Cập nhật/Thêm styles cho các danh mục (Table of Contents & Table of Figures)
    from docx.enum.style import WD_STYLE_TYPE
    toc_styles = [
        'TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5', 
        'TOC1', 'TOC2', 'TOC3', 'TOC4', 'TOC5',
        'Table of Figures', 'TableofFigures', 
        'Table of Tables', 'TableofTables',
        'Index', 'Index 1', 'Index 2'
    ]
    style_id_map = {
        'TOC 1': 'toc1', 'TOC 2': 'toc2', 'TOC 3': 'toc3', 'TOC 4': 'toc4', 'TOC 5': 'toc5',
        'TOC1': 'toc1', 'TOC2': 'toc2', 'TOC3': 'toc3', 'TOC4': 'toc4', 'TOC5': 'toc5',
        'Table of Figures': 'tableoffigures', 'TableofFigures': 'tableoffigures',
        'Table of Tables': 'tableoffigures', 'TableofTables': 'tableoffigures'
    }
    for style_name in toc_styles:
        try:
            if style_name not in doc.styles:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = doc.styles[style_name]
                
            # Áp dụng ID chuẩn của Word cho các built-in styles để tránh lỗi Word nhận diện nhầm làm custom style
            if style_name in style_id_map:
                style_el = style.element
                style_el.set(qn('w:styleId'), style_id_map[style_name])
                if qn('w:customStyle') in style_el.attrib:
                    del style_el.attrib[qn('w:customStyle')]

            style.font.name = font_name
            style.font.size = Pt(14)
            style.font.bold = False
            style.font.italic = False
            
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:cs'), font_name)
            
            # Thiết lập rõ ràng w:val="0" cho các thuộc tính bold/italic trong XML để ghi đè mọi kế thừa từ mẫu Word gốc
            for tag in ['w:b', 'w:bCs', 'w:i', 'w:iCs']:
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn('w:val'), '0')
            
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.first_line_indent = Pt(0)
            
            # Thiết lập lùi lề phân cấp cho các cấp độ mục lục (TOC 1-5) để hiển thị lùi lề hình cây đẹp mắt
            indent_pt = 0
            if 'TOC' in style_name:
                try:
                    level = int(re.search(r'\d+', style_name).group())
                    if level > 1:
                        # Cấp 2 lùi 18pt (0.63cm), Cấp 3 lùi 36pt (1.27cm), v.v.
                        indent_pt = (level - 1) * 18
                except Exception as e:
                        logging.warning(f"Silent error ignored: {e}")
            style.paragraph_format.left_indent = Pt(indent_pt)
            style.paragraph_format.right_indent = Pt(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            set_contextual_spacing(style, True)
        except Exception as e:
            print(f"Lỗi cấu hình style {style_name}:", e)

    # Cấu hình styles cho các Heading (Heading 1 đến Heading 9) trực tiếp trong Stylesheet
    for lvl in range(1, 10):
        style_name = f"Heading {lvl}"
        try:
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_style.font.name = font_name
                
                if lvl == 1:
                    h_size = float(opts.get('heading1_size', 14))
                    h_bold = opts.get('heading1_bold', True)
                    h_italic = False
                elif lvl == 2:
                    h_size = float(opts.get('heading2_size', 14))
                    h_bold = opts.get('heading2_bold', True)
                    h_italic = False
                elif lvl == 3:
                    h_size = float(opts.get('heading3_size', 13))
                    h_bold = True
                    h_italic = opts.get('heading3_italic', False)
                else:
                    # Heading 4+
                    h_size = float(opts.get('heading3_size', 13))
                    h_bold = True
                    h_italic = opts.get('heading3_italic', False)
                    
                h_style.font.size = Pt(h_size)
                h_style.font.bold = h_bold
                h_style.font.italic = h_italic
                
                # Áp dụng trực tiếp vào XML của style để ghi đè triệt để kế thừa cũ
                rPr = h_style.element.get_or_add_rPr()
                
                b_val = '1' if h_bold else '0'
                for tag in ['w:b', 'w:bCs']:
                    el = rPr.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        rPr.append(el)
                    el.set(qn('w:val'), b_val)
                    
                i_val = '1' if h_italic else '0'
                for tag in ['w:i', 'w:iCs']:
                    el = rPr.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        rPr.append(el)
                    el.set(qn('w:val'), i_val)
        except Exception as e:
            print(f"Lỗi cấu hình style {style_name}:", e)

    # Cấu hình style Caption trực tiếp trong Stylesheet
    try:
        if 'Caption' in doc.styles:
            cap_style = doc.styles['Caption']
            cap_style.font.name = font_name
            cap_style.font.size = Pt(14)
            cap_style.font.bold = False
            cap_style.font.italic = True
            
            rPr = cap_style.element.get_or_add_rPr()
            for tag in ['w:b', 'w:bCs']:
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn('w:val'), '0')
            for tag in ['w:i', 'w:iCs']:
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn('w:val'), '1')
    except Exception as e:
        print("Lỗi cấu hình style Caption:", e)

    auto_num = opts.get('auto_number_headings', False)
    heading_counters = [0] * 9
    fig_cnt, tbl_cnt = 0, 0
    do_admin = opts.get('format_admin_parts', False)
    in_recipients = False
    in_front_matter_directory = False
    current_h1_is_no_num = False

    # --- 3. XỬ LÝ TỪNG PARAGRAPH ---
    para_section_type = {}
    try:
        parts = partition_paragraphs_by_section(doc)
        section_types = classify_sections(doc)
        for idx, part in enumerate(parts):
            sect_type = section_types[idx] if idx < len(section_types) else "body"
            for p_in_part in part:
                para_section_type[p_in_part._p] = sect_type
    except Exception as e:
        print("Lỗi phân nhóm section cho paragraph loop:", e)

    cover_paras_set = {p._p for p in cover_paras} if cover_table is None else set()
    for p in doc.paragraphs:
        if not format_cover and para_section_type.get(p._p) == "cover":
            continue
        if p._p in cover_paras_set:
            continue
        # Nếu có cover_table, bỏ qua không định dạng các paragraph trống thuộc section cover
        if cover_table is not None and para_section_type.get(p._p) == "cover":
            continue
        # Xóa toàn bộ ngắt trang thủ công (br type="page") trong XML của paragraph
        for br in p._p.xpath('.//*[local-name()="br" and @*[local-name()="type"]="page"]'):
            try:
                br.getparent().remove(br)
            except Exception as e:
                    logging.warning(f"Silent error ignored: {e}")
                    
        text = p.text.strip()
        style_name = p.style.name if p.style else 'Normal'

        if not text and not paragraph_has_image(p):
            if style_name != 'TOC Heading':
                continue

        # Kích hoạt không cách giữa các đoạn cùng style
        try:
            set_contextual_spacing(p, contextual_spacing)
        except Exception as e:
                logging.warning(f"Silent error ignored: {e}")

        # Clean "Chương X: " or "Chương X :" to "Chương X " (removing the colon)
        cleaned_text = re.sub(r'^(CHƯƠNG\s+[IVX\d]+)\s*:\s*', r'\1 ', text, flags=re.IGNORECASE)
        if cleaned_text != text:
            p.text = cleaned_text
            text = cleaned_text

        text_upper = text.upper()

        # Kiểm tra xem có phải tiêu đề danh mục/mục lục hay không
        text_lower = text.lower()
        is_directory_title = False
        if len(text) < 100:
            if "mục lục" in text_lower or "muc luc" in text_lower:
                is_directory_title = True
            elif len(text) < 60 and (text_lower.startswith("danh mục") or text_lower.startswith("danh muc")):
                if any(x in text_lower for x in ["hình", "hinh", "bảng", "bang", "ký hiệu", "ky hieu", "viết tắt", "viet tat", "thuật ngữ", "thuat ngu"]):
                    is_directory_title = True

        if is_directory_title:
            in_front_matter_directory = True

        # Kiểm tra Mục lục
        is_muc_luc = (text_upper == "MỤC LỤC" 
                      or text_upper.startswith("MỤC LỤC ") 
                      or re.sub(r'[^a-z]', '', text.lower()) == "mucluc"
                      or style_name == 'TOC Heading')

        if is_muc_luc:
            in_front_matter_directory = True
            if not text:
                p.text = "MỤC LỤC"
                text = "MỤC LỤC"
                text_upper = "MỤC LỤC"
            # MỤC LỤC: Không gán Heading, chỉ định dạng Normal, bold, centered, 14pt (hoặc cỡ chữ h1)
            safe_set_style(doc, p, 'Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.page_break_before = True
            
            # Cập nhật font chữ cho các runs
            for r in p.runs:
                set_run_font(r, font_name, float(opts.get('heading1_size', 14)), bold=True, italic=False)
            continue

        # --- XỬ LÝ HÌNH ẢNH & BẢNG BIỂU ---
        # Căn giữa hình ảnh (sử dụng hàm kiểm tra namespace-agnostic)
        if paragraph_has_image(p) and not text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.keep_with_next = True  # Giữ ảnh luôn cùng trang với dòng chú thích bên dưới
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_contextual_spacing(p, True)
            
            # Căn giữa cả ảnh dạng nổi (floating anchor) nếu có
            center_floating_images(p)
            
            # Xóa các khoảng trắng thừa/dấu tab làm lệch lề ảnh
            for r in p.runs:
                if r.text:
                    r.text = r.text.strip(' \t\n\r ')
                    
            # Xóa các thẻ w:tab trong XML của ảnh
            for tab in p._p.xpath('.//*[local-name()="tab"]'):
                try:
                    tab.getparent().remove(tab)
                except Exception as e:
                        logging.warning(f"Silent error ignored: {e}")
            continue

        # Kiểm tra và định dạng chú thích Hình/Ảnh/Bảng tự động (tương thích danh mục động SEQ)
        is_fig_caption = (not in_front_matter_directory) and (not is_directory_line(text)) and (len(text) < 200 and re.match(r'^(Hình|Ảnh)\b', text, re.IGNORECASE))
        is_tbl_caption = (not in_front_matter_directory) and (not is_directory_line(text)) and (len(text) < 200 and re.match(r'^Bảng\b', text, re.IGNORECASE))

        if is_fig_caption:
            fig_cnt += 1
            clean = re.sub(r'^(Hình|Ảnh)\s*[\d\.\-\s:]*', '', text, flags=re.IGNORECASE).strip().lstrip(':. ')
            effective_h1 = heading_counters[0] if heading_counters[0] > 0 else 1
            make_caption_paragraph(p, "Hình", effective_h1, fig_cnt, clean, font_name, 14)
            safe_set_style(doc, p, 'Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            p.paragraph_format.right_indent = None
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_contextual_spacing(p, True)
            continue

        if is_tbl_caption:
            tbl_cnt += 1
            clean = re.sub(r'^Bảng\s*[\d\.\-\s:]*', '', text, flags=re.IGNORECASE).strip().lstrip(':. ')
            effective_h1 = heading_counters[0] if heading_counters[0] > 0 else 1
            make_caption_paragraph(p, "Bảng", effective_h1, tbl_cnt, clean, font_name, 14)
            safe_set_style(doc, p, 'Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            p.paragraph_format.right_indent = None
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_contextual_spacing(p, True)
            p.paragraph_format.keep_with_next = True  # Giữ tên bảng biểu luôn cùng trang với bảng bên dưới
            continue

        # Định nghĩa các từ khóa Heading 1 không đánh số
        h1_no_num_keywords = [
            "lời cảm ơn", "lời cam đoan", "danh mục hình ảnh", "danh mục bảng biểu",
            "danh mục các thuật ngữ viết tắt", "danh mục thuật ngữ viết tắt",
            "danh mục các thuật ngữ, ký hiệu và các chữ viết tắt",
            "danh mục thuật ngữ, ký hiệu và từ viết tắt",
            "danh mục các thuật ngữ, ký hiệu và từ viết tắt",
            "mở đầu", "phần mở đầu", "lời mở đầu", "lời nói đầu",
            "kết luận", "tài liệu tham khảo", "phụ lục",
            "loi cam on", "loi cam doan", "danh muc hinh anh", "danh muc bang bieu",
            "danh muc cac thuat ngu viet tat", "danh muc thuat ngu viet tat",
            "danh muc cac thuat ngu, ky hieu va cac chu viet tat",
            "danh muc thuat ngu, ky hieu va tu viet tat",
            "danh muc cac thuat ngu, ky hieu va tu viet tat",
            "mo dau", "phan mo dau", "loi mo dau", "loi noi dau",
            "ket luan", "tai lieu tham khao", "phu luoc", "phu luc"
        ]
        
        # Hàm kiểm tra xem có phải Heading 1 không đánh số hay không (phải là dòng ngắn và khớp với từ khóa)
        text_lower_clean = re.sub(r'[^a-z0-9àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]', '', text.lower()).strip()
        h1_no_num_clean_keywords = [
            re.sub(r'[^a-z0-9àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]', '', k)
            for k in h1_no_num_keywords
        ]
        is_h1_no_num = (len(text) < 80) and (text_lower_clean in h1_no_num_clean_keywords or is_directory_title)

        is_h1, is_h2, is_h3 = False, False, False
        heading_level = None
        if not is_directory_line(text):
            explicit_outline = None
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                outlineLvl = pPr.find(qn('w:outlineLvl'))
                if outlineLvl is not None:
                    try:
                        explicit_outline = int(outlineLvl.get(qn('w:val')))
                    except (ValueError, TypeError):
                        pass

            match_h_num = re.match(r'^(\d+(?:\.\d+)+)(?:\D|$)', text)
            if match_h_num:
                num_str = match_h_num.group(1)
                level = num_str.count('.') + 1
                if level == 2:
                    heading_level = 2
                elif level >= 3:
                    heading_level = level
            elif is_h1_no_num or re.match(r'^CHƯƠNG\s+[IVX\d]+', text, re.IGNORECASE) or re.match(r'^I{1,3}\.\s', text):
                heading_level = 1
            elif explicit_outline is not None:
                heading_level = explicit_outline + 1
            else:
                if style_name == 'Heading 1':
                    heading_level = 1
                elif style_name == 'Heading 2':
                    heading_level = 2
                elif style_name.startswith('Heading '):
                    try:
                        lvl = int(style_name.split(' ')[-1])
                        if lvl >= 3:
                            heading_level = lvl
                    except ValueError:
                        pass
            # --- KHỬ NHẬN DIỆN NHẦM HEADING (DECLASSIFICATION RULES) ---
            # Luật 2: Nếu nhận diện là Heading bằng từ khóa/pattern số nhưng style thực tế là Normal/Body và quá dài (> 100 kí tự)
            if heading_level is not None and not style_name.startswith('Heading'):
                if len(text) > 100:
                    heading_level = None

            # Luật 1: Nếu style thực tế là Heading nhưng text không khớp pattern số heading và quá dài (> 150 kí tự)
            if heading_level is not None and style_name.startswith('Heading'):
                has_pattern = re.match(r'^(\d+(?:\.\d+)+|CHƯƠNG\s+[IVX\d]+|I{1,3}\.)\b', text, re.IGNORECASE) is not None
                if not has_pattern and len(text) > 150:
                    heading_level = None

            # Luật 3: Khử Heading 1 nhận diện nhầm (không có ký tự bắt đầu của Chương và không nằm trong H1 không đánh số)
            if heading_level == 1 and not is_h1_no_num:
                has_chapter_pattern = re.match(r'^(CHƯƠNG|Chương|[IVX\d]+)\b', text, re.IGNORECASE) is not None
                if not has_chapter_pattern:
                    heading_level = 2
            
            # Luật 4: Khử nhận diện nhầm tiêu đề chương trong phần bố cục/tóm tắt (nằm trong phần không đánh số như MỞ ĐẦU)
            if heading_level == 1 and not is_h1_no_num and not style_name.startswith('Heading'):
                if current_h1_is_no_num:
                    ch_num = get_chapter_number(text)
                    if ch_num is not None:
                        expected_ch = heading_counters[0] + 1
                        if ch_num > expected_ch or any(c.islower() for c in text):
                            heading_level = None
            
            if heading_level == 1:
                is_h1 = True
            elif heading_level == 2:
                is_h2 = True
            elif heading_level is not None and heading_level >= 3:
                is_h3 = True

        # Xác định xem có thoát khỏi danh mục đầu trang hay không
        is_real_heading = (is_h1 or is_h2 or is_h3) and (not is_directory_title)
        if not is_real_heading:
            if re.match(r'^CHƯƠNG\s+[IVX\d]+', text, re.IGNORECASE):
                is_real_heading = True
            elif text_upper in ["MỞ ĐẦU", "PHẦN MỞ ĐẦU", "LỜI MỞ ĐẦU", "LỜI NÓI ĐẦU"]:
                is_real_heading = True

        if is_real_heading and not is_directory_line(text):
            in_front_matter_directory = False

        # --- XỬ LÝ HÀNH CHÍNH (Nghị định 30) ---
        admin_handled = False
        if do_admin:
            # Quốc hiệu
            if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text_upper:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, font_name, body_size - 1, bold=True, italic=False)
                    r.text = r.text.upper()
                admin_handled = True

            # Tiêu ngữ
            elif "ĐỘC LẬP" in text_upper and ("TỰ DO" in text_upper or "TỰ DO" in text_upper) and "HẠNH PHÚC" in text_upper:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, font_name, body_size, bold=True, italic=False)
                admin_handled = True

            # Địa danh, ngày tháng
            elif re.search(r'ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+', text, re.IGNORECASE):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, font_name, body_size, bold=False, italic=True)
                admin_handled = True

            # Nơi nhận
            elif text.lower().startswith("nơi nhận:"):
                in_recipients = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, font_name, 12, bold=True, italic=True)
                admin_handled = True

            elif in_recipients and (text.startswith("-") or text.startswith("+") or len(text) < 30):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, font_name, 11, bold=False, italic=False)
                if not text.startswith("-") and len(text) > 40:
                    in_recipients = False
                else:
                    admin_handled = True
            else:
                in_recipients = False

        if admin_handled:
            continue

        # --- HEADINGS ---
        if is_h1:
            safe_set_style(doc, p, 'Heading 1')
            current_h1_is_no_num = is_h1_no_num

            if not is_h1_no_num:
                heading_counters[0] += 1
                for idx in range(1, 9):
                    heading_counters[idx] = 0
                fig_cnt, tbl_cnt = 0, 0

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.page_break_before = True

            h1_size = float(opts.get('heading1_size', 14))
            h1_bold = opts.get('heading1_bold', True)
            h1_upper = opts.get('heading1_uppercase', True)

            if auto_num:
                if is_h1_no_num:
                    p.text = text
                else:
                    has_newline = '\n' in text
                    clean = re.sub(r'^(CHƯƠNG\s+[IVX\d]+[\.:]?|I{1,3}\.|[\d\.]+)\s*', '', text, flags=re.IGNORECASE).strip().lstrip(':. ')
                    if has_newline:
                        p.text = f"CHƯƠNG {heading_counters[0]}.\n{clean}"
                    else:
                        p.text = f"CHƯƠNG {heading_counters[0]}. {clean}"
            
            if h1_upper:
                p.text = p.text.upper()

            for r in p.runs:
                set_run_font(r, font_name, h1_size)
                r.bold = None
                r.italic = None
                try: r.font.cs_bold = None
                except Exception: pass
                try: r.font.cs_italic = None
                except Exception: pass

        elif is_h2:
            safe_set_style(doc, p, 'Heading 2')

            heading_counters[1] += 1
            for idx in range(2, 9):
                heading_counters[idx] = 0

            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)

            h2_size = float(opts.get('heading2_size', 14))
            h2_bold = opts.get('heading2_bold', True)

            if auto_num:
                clean = re.sub(r'^[\d\.\-\s:]+', '', text).strip().lstrip(':. ')
                effective_h1 = heading_counters[0] if heading_counters[0] > 0 else 1
                p.text = f"{effective_h1}.{heading_counters[1]}. {clean}"
            for r in p.runs:
                set_run_font(r, font_name, h2_size)
                r.bold = None
                r.italic = None
                try: r.font.cs_bold = None
                except Exception: pass
                try: r.font.cs_italic = None
                except Exception: pass

        elif is_h3:
            level = heading_level if heading_level is not None else 3
            target_style = f"Heading {level}"
            safe_set_style(doc, p, target_style if target_style in doc.styles else 'Heading 3')

            if level <= 9:
                heading_counters[level - 1] += 1
                for idx in range(level, 9):
                    heading_counters[idx] = 0

            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)

            h3_size = float(opts.get('heading3_size', 13))
            h3_italic = opts.get('heading3_italic', False)

            if auto_num:
                clean = re.sub(r'^([\d\.\-\s:]+|[a-z]\)\s*)', '', text, flags=re.IGNORECASE).strip().lstrip(':. ')
                prefix_parts = []
                for i in range(level):
                    val = heading_counters[i]
                    if val == 0 and i < 2:
                        val = 1
                    prefix_parts.append(str(val))
                p.text = f"{'.'.join(prefix_parts)}. {clean}"
            for r in p.runs:
                set_run_font(r, font_name, h3_size)
                r.bold = None
                r.italic = None
                try: r.font.cs_bold = None
                except Exception: pass
                try: r.font.cs_italic = None
                except Exception: pass

        else:
            if (in_front_matter_directory or is_toc_or_directory_paragraph(p)) and not is_directory_title and not is_muc_luc:
                # Định dạng danh mục theo yêu cầu: cỡ chữ 14 thường, dãn 1.5, 0pt 0pt, không in đậm
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                # Để thừa kế lùi lề phân cấp từ stylesheet cho các mục lục TOC để giữ cấu trúc hình cây
                style_name_upper = (p.style.name or '').upper()
                if style_name_upper.startswith('TOC'):
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.first_line_indent = None
                else:
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.first_line_indent = Pt(0)
                    
                p.paragraph_format.right_indent = Pt(0)
                try:
                    set_contextual_spacing(p, True)
                except Exception as e:
                        logging.warning(f"Silent error ignored: {e}")
                for r in p.runs:
                    set_run_font(r, font_name, 14, bold=False, italic=False)
                    # Thiết lập rõ ràng w:val="0" cho các thuộc tính bold/italic trong XML để ghi đè mọi kế thừa từ mẫu Word gốc
                    rPr = r.element.get_or_add_rPr()
                    for tag in ['w:b', 'w:bCs', 'w:i', 'w:iCs']:
                        el = rPr.find(qn(tag))
                        if el is None:
                            el = OxmlElement(tag)
                            rPr.append(el)
                        el.set(qn('w:val'), '0')
            else:
                p.alignment = body_alignment
                p.paragraph_format.line_spacing = line_spacing
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.left_indent = None
                p.paragraph_format.right_indent = None
                
                is_references = para_section_type.get(p._p) == "references"
                
                # Check if it starts with numbered item (e.g. 1., 2.)
                is_list_item = False
                if re.match(r'^\d+\.', text):
                    is_list_item = True
                    
                if is_list_item or is_references:
                    p.paragraph_format.first_line_indent = Pt(0)
                else:
                    p.paragraph_format.first_line_indent = first_line_indent
                
                # Xóa khoảng trắng/tab thủ công ở đầu đoạn văn để Word tự thụt lề bằng Paragraph Format
                if p.runs:
                    for r in p.runs:
                        if r.text:
                            r.text = r.text.lstrip(' \t\n\r ')
                            if r.text:
                                break

                for r in p.runs:
                    set_run_font(r, font_name, body_size)
                try:
                    set_contextual_spacing(p, contextual_spacing)
                except Exception as e:
                        logging.warning(f"Silent error ignored: {e}")

    # --- 4. BẢNG BIỂU (Định dạng đệ quy bao gồm cả bảng lồng bên trong) ---
    def format_table_recursive(tbl):
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in tbl.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph_has_image(paragraph):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.right_indent = Pt(0)
                        
                        # Xóa các thẻ w:tab trong XML của ảnh trong ô bảng
                        for tab in paragraph._p.xpath('.//*[local-name()="tab"]'):
                            try:
                                tab.getparent().remove(tab)
                            except Exception as e:
                                    logging.warning(f"Silent error ignored: {e}")
                                
                        # Dọn dẹp khoảng trắng/tab của run
                        for r in paragraph.runs:
                            if r.text:
                                r.text = r.text.strip(' \t\n\r ')
                    else:
                        # Chỉ thay đổi căn lề nếu ban đầu không phải là Center hoặc Right để bảo vệ tiêu đề/căn lề đặc biệt
                        if paragraph.alignment not in [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.right_indent = Pt(0)
                        
                        # Dọn dẹp khoảng trắng/tab ở đầu đoạn văn trong bảng
                        if paragraph.runs:
                            for r in paragraph.runs:
                                if r.text:
                                    r.text = r.text.lstrip(' \t\n\r ')
                                    if r.text:
                                        break
                                        
                    paragraph.paragraph_format.line_spacing = line_spacing
                    paragraph.paragraph_format.space_before = Pt(space_before)
                    paragraph.paragraph_format.space_after = Pt(space_after)
                    try:
                        set_contextual_spacing(paragraph, contextual_spacing)
                    except Exception as e:
                            logging.warning(f"Silent error ignored: {e}")
                    for r in paragraph.runs:
                        set_run_font(r, font_name, body_size)
                
                # Gọi đệ quy cho các bảng lồng bên trong cell
                for nested_tbl in cell.tables:
                    format_table_recursive(nested_tbl)

    # Phân loại section cho từng table
    table_section_types = {}
    try:
        element_section_indices = get_element_section_types(doc)
        for table in doc.tables:
            tbl_el = table._tbl
            sec_idx = element_section_indices.get(tbl_el, 0)
            sec_type = section_types[sec_idx] if sec_idx < len(section_types) else "body"
            table_section_types[table] = sec_type
    except Exception as e:
        print("Lỗi khi xác định section cho table:", e)

    for table in doc.tables:
        if not format_cover and table_section_types.get(table) == "cover":
            continue
        if cover_table is not None and table == cover_table:
            continue
        format_table_recursive(table)

    # --- 5. ĐÁNH SỐ TRANG ---
    if opts.get('add_page_numbers', True):
        section_types = classify_sections(doc)
        for idx, section in enumerate(doc.sections):
            if not format_cover and idx in cover_section_indices:
                continue
            sect_type = section_types[idx] if idx < len(section_types) else "body"
            
            # Đảm bảo hủy liên kết (unlink) để tránh đè hoặc ảnh hưởng tới section khác
            section.footer.is_linked_to_previous = False
            section.first_page_footer.is_linked_to_previous = False
            section.header.is_linked_to_previous = False
            section.first_page_header.is_linked_to_previous = False
            
            # Xóa sạch toàn bộ nội dung cũ trong tất cả các header/footer
            clear_header_footer(section.footer)
            clear_header_footer(section.first_page_footer)
            clear_header_footer(section.header)
            clear_header_footer(section.first_page_header)
            
            if sect_type == "cover":
                section.different_first_page_header_footer = True
            else:
                section.different_first_page_header_footer = False
                if sect_type in ["thanks", "list", "body"]:
                    header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
                    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header_p.paragraph_format.space_before = Pt(0)
                    header_p.paragraph_format.space_after = Pt(0)
                    header_p.paragraph_format.line_spacing = 1.0
                    header_p.paragraph_format.first_line_indent = Pt(0)
                    header_p.paragraph_format.left_indent = Pt(0)
                    header_p.paragraph_format.right_indent = Pt(0)
                    run = header_p.add_run()
                    set_run_font(run, font_name, body_size, bold=False, italic=False)
                    add_page_number_field(run)
            
            # Cấu hình w:pgNumType và start bắt đầu
            if sect_type == "thanks":
                set_section_page_numbering(section, fmt="lowerRoman", start=1)
            elif sect_type == "toc":
                set_section_page_numbering(section, fmt="lowerRoman", start=None)
            elif sect_type == "list":
                # Bắt đầu đánh số trang từ số La mã ii (2) cho danh mục 
                # (để bù lại trang Mục lục không được đánh số và không hiển thị, giúp Lời cảm ơn là i, Danh mục là ii)
                set_section_page_numbering(section, fmt="lowerRoman", start=2)
            elif sect_type == "body":
                set_section_page_numbering(section, fmt="decimal", start=1)
            elif sect_type == "references":
                set_section_page_numbering(section, fmt="decimal", start=None)

    # --- TỰ ĐỘNG CHÈN TRƯỜNG DANH MỤC HÌNH ẢNH, BẢNG BIỂU & MỤC LỤC NẾU THIẾU ---
    for p in list(doc.paragraphs):
        text_upper = p.text.strip().upper()
        clean_text = re.sub(r'\s+', ' ', text_upper)
        if clean_text in ["DANH MỤC HÌNH ẢNH", "DANH MỤC HÌNH"]:
            parent = p._p.getparent()
            p_index = parent.index(p._p)
            if not parent_already_has_toc(parent, p_index, "Hình"):
                insert_toc_field_after(p, doc, ' TOC \\h \\z \\c "Hình" ', "TableofFigures", font_name)
        elif clean_text in ["DANH MỤC BẢNG BIỂU", "DANH MỤC BẢNG"]:
            parent = p._p.getparent()
            p_index = parent.index(p._p)
            if not parent_already_has_toc(parent, p_index, "Bảng"):
                insert_toc_field_after(p, doc, ' TOC \\h \\z \\c "Bảng" ', "TableofFigures", font_name)
        elif clean_text == "MỤC LỤC":
            parent = p._p.getparent()
            p_index = parent.index(p._p)
            if not parent_already_has_toc(parent, p_index, "\\o"):
                insert_toc_field_after(p, doc, ' TOC \\o "1-4" \\h \\z \\u ', "TOC1", font_name)

    # --- DỌN DẸP PARAGRAPH RỖNG GIỮA TIÊU ĐỀ DANH MỤC VÀ NỘI DUNG TOC ---
    try:
        from lxml import etree as _et
        _xpath_instr = _et.XPath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        _all_paras = list(doc.paragraphs)
        _paras_to_del = []
        for _i, _p in enumerate(_all_paras):
            _tu = _p.text.strip().upper()
            _ct = re.sub(r'\s+', ' ', _tu)
            if _ct in ['DANH MỤC HÌNH ẢNH', 'DANH MỤC HÌNH', 'DANH MỤC BẢNG BIỂU', 'DANH MỤC BẢNG', 'MỤC LỤC',
                        'DANH MỤC CÁC THUẬT NGỮ VIẾT TẮT', 'DANH MỤC THUẬT NGỮ VIẾT TẮT',
                        'DANH MỤC CÁC THUẬT NGỮ, KÝ HIỆU VÀ CÁC CHỮ VIẾT TẮT',
                        'DANH MỤC THUẬT NGỮ, KÝ HIỆU VÀ TỪ VIẾT TẮT',
                        'DANH MỤC CÁC THUẬT NGỮ, KÝ HIỆU VÀ TỪ VIẾT TẮT']:
                _j = _i + 1
                while _j < len(_all_paras):
                    _np = _all_paras[_j]
                    _nt = _np.text.strip()
                    _has_toc = len(_xpath_instr(_np._p)) > 0
                    _has_img = paragraph_has_image(_np)
                    _pPr = _np._p.find(qn('w:pPr'))
                    _has_sect = _pPr is not None and _pPr.find(qn('w:sectPr')) is not None
                    if _nt or _has_toc or _has_img:
                        break  # Gặp nội dung thực, dừng
                    if _has_sect:
                        _j += 1
                        continue  # Giữ paragraph mang section break
                    _paras_to_del.append(_np)
                    _j += 1
        for _p in _paras_to_del:
            try:
                _p._p.getparent().remove(_p._p)
            except Exception as e:
                    logging.warning(f"Silent error ignored: {e}")
    except Exception as _e:
        print('Lỗi dọn dẹp paragraph rỗng giữa danh mục:', _e)

    # --- TỰ ĐỘNG CẬP NHẬT MỤC LỤC & DANH MỤC KHI MỞ FILE ---
    try:
        settings_el = doc.settings.element if hasattr(doc.settings, 'element') else doc.settings._element
        update_fields = settings_el.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            insert_element_in_order(settings_el, update_fields, SETTINGS_ORDER)
    except Exception as e:
        print("Lỗi cấu hình tự động cập nhật danh mục:", e)

    # --- LOẠI BỎ TRANG TRỐNG VÀ DỌN DẸP ---
    # Loại bỏ page_break_before ở các đoạn bắt đầu section mới (đã có section break) để tránh sinh ra 2 trang trắng liên tiếp
    # Section break đã đảm bảo paragraph bắt đầu ở trang mới, nên page_break_before phải tắt để tránh tạo trang trắng đôi
    if opts.get('add_page_numbers', True):
        # 1. Tắt page_break_before cho trigger_paras
        for p in trigger_paras:
            p.paragraph_format.page_break_before = False
            
        # 2. Tắt page_break_before cho bất kỳ paragraph nào đứng ngay sau một section break (đầu của một section mới)
        is_first_of_section = True
        for p in doc.paragraphs:
            if is_first_of_section:
                p.paragraph_format.page_break_before = False
                is_first_of_section = False
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                is_first_of_section = True
            
    # Loại bỏ các đoạn trống thừa ở cuối tài liệu (tránh tạo trang trắng ở cuối)
    for p in reversed(doc.paragraphs):
        if not p.text.strip():
            p_el = p._p
            p_el.getparent().remove(p_el)
        else:
            break

    # --- CẬP NHẬT CẤP ĐỘ MỤC LỤC LÊN HEADING 4 CHO TRƯỜNG HIỆN CÓ ---
    try:
        from lxml import etree as _et
        _xpath_instr = _et.XPath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for instr in _xpath_instr(doc.element.body):
            if instr.text and 'TOC' in instr.text:
                new_text = re.sub(r'\\o\s+["\']1-[23]["\']', '\\\\o "1-4"', instr.text)
                if new_text != instr.text:
                    instr.text = new_text
    except Exception as e:
        print("Lỗi khi cập nhật cấp độ mục lục lên Heading 4:", e)

    # --- LƯU KẾT QUẢ ---
    doc.save(output_path)
