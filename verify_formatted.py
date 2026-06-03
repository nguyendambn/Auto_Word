import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

output_path = r"c:\Users\The Linh\Desktop\2022604488_MaiXuanBac_formatted_test_2.docx"
if not os.path.exists(output_path):
    print("File output chưa tồn tại!")
    sys.exit(1)

doc = Document(output_path)
print("--- THÔNG TIN FILE SAU KHI FORMAT ---")
print(f"Tổng số paragraph: {len(doc.paragraphs)}")
print(f"Tổng số table: {len(doc.tables)}")
print(f"Tổng số section: {len(doc.sections)}")

print("\n--- CÁC SECTION VÀ TRANG BẮT ĐẦU ---")
for idx, sec in enumerate(doc.sections):
    print(f"Section {idx}: page_width={sec.page_width.mm:.1f}mm, page_height={sec.page_height.mm:.1f}mm")
    print(f"  left_margin={sec.left_margin.mm:.1f}mm, right_margin={sec.right_margin.mm:.1f}mm")
    print(f"  top_margin={sec.top_margin.mm:.1f}mm, bottom_margin={sec.bottom_margin.mm:.1f}mm")
    print(f"  different_first_page={sec.different_first_page_header_footer}")
    # Inspect footer/header paragraphs
    header_text = "".join(p.text for p in sec.header.paragraphs).strip()
    footer_text = "".join(p.text for p in sec.footer.paragraphs).strip()
    print(f"  Header text: {repr(header_text)}")
    print(f"  Footer text: {repr(footer_text)}")

print("\n--- 40 PARAGRAPH ĐẦU TIÊN ---")
for i, p in enumerate(doc.paragraphs[:40]):
    print(f"[{i:02d}][Style: {p.style.name}]: {repr(p.text)}")
