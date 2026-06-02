import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:\Users\The Linh\Desktop\Auto-Word\_verify_headings.docx')
print(f"Index 94: Style={doc.paragraphs[94].style.name} | Text={repr(doc.paragraphs[94].text)}")
