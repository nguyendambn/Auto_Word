import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:\Users\The Linh\Desktop\CNTT_2022607319_NguyenVanDam_Baocao.docx')
for i in range(160, 175):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"Index {i} | Style: {p.style.name} | Text: {repr(p.text)}")
