import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:\Users\The Linh\Desktop\CNTT_2022607319_NguyenVanDam_Baocao.docx')

print("--- HEADERS AND FOOTERS IN ORIGINAL DOCUMENT ---")
for idx, s in enumerate(doc.sections):
    print(f"\nSection {idx}:")
    print(f"  different_first_page_header_footer: {s.different_first_page_header_footer}")
    
    # Header
    print(f"  Header runs: {[r.text for p in s.header.paragraphs for r in p.runs]}")
    # First Page Header
    print(f"  First Page Header runs: {[r.text for p in s.first_page_header.paragraphs for r in p.runs]}")
    # Footer
    print(f"  Footer runs: {[r.text for p in s.footer.paragraphs for r in p.runs]}")
    # First Page Footer
    print(f"  First Page Footer runs: {[r.text for p in s.first_page_footer.paragraphs for r in p.runs]}")
