import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:\Users\The Linh\Desktop\CNTT_2022607319_NguyenVanDam_Baocao.docx')

print("--- ALL HEADERS AND FOOTERS IN ORIGINAL DOCUMENT ---")
for idx, s in enumerate(doc.sections):
    print(f"\nSection {idx}:")
    for name in ['header', 'first_page_header', 'even_page_header', 'footer', 'first_page_footer', 'even_page_footer']:
        hf = getattr(s, name)
        p_texts = [p.text for p in hf.paragraphs]
        p_xmls = [docx.oxml.xmlchemy.serialize_for_reading(p._p) for p in hf.paragraphs if p.text.strip()]
        if p_texts or p_xmls:
            print(f"  {name}:")
            print(f"    Paragraph texts: {p_texts}")
            if p_xmls:
                print(f"    Paragraph XMLs with text: {p_xmls}")
